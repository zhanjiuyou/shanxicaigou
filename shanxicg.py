# 2020.9.13日
# 本次爬虫主要是实现陕西政府采购网的爬虫
# 该爬虫以技术交流为准，切勿用于任何非法用途，后果自负
# 我的个人博客：https://jiaokangyang.com

from selenium import webdriver
import pymysql
import time
from pyquery import PyQuery as pq
from docx import *

# 由于该网站是javascript来异步加载的，而且requests不能正常获取，这里我们使用selenium爬取

url = 'http://ccgp-shaanxi.gov.cn/notice/list.do?noticetype=3&province=province'
shuru = input('请输入要爬取的区域名称，确保在网站的范围内：')
shuru2 = input('请输入上述城市中要筛选的区域名：\n(如不需要筛选则直接敲击回车键开始抓取)\n')
# 创建一个空文档，用于后面的文档保存
document = Document()

# 链接数据库
conn = pymysql.connect(host='localhost',user='root',password='123456',port=3306,database='shanxi')
# 获取游标，后面进行sql语句的执行
cursor = conn.cursor()
#  如果表不存在则创建一个名为shuju的表。这里切记mysql中的列名字不需要加双引号。
cursor.execute("create table if not exists shuju(id int not null auto_increment primary key,title varchar(50) not null,url varchar(100) not null,riqi varchar(20) not null);")


# 打开谷歌浏览器
brower = webdriver.Chrome()
# 打开采购信息的网页

brower.get(url)
# 打开网页后，点击对应城市的标签，然后异步加载的内容进行加载。
if shuru != '':
    brower.find_element_by_link_text(shuru).click()
# 这块由于代码自动操作太快，有时出现内容没更新的情况，所以我们等待两秒
time.sleep(2)


# 该函数完成单页内容的采集输出
def get_onepage(html):
    html = pq(html)
    a = '[' + shuru2 + ']'
    lists = html('.list-box table tbody tr').items()
    for list in lists:
        if shuru2 != '':
            b = list('td:nth-child(2)').text() # 使用pyquery的伪类用法查找第二个元素内的名字
            if b == a:  # 对比分析，如果和我们输入的区域名字相同，则打印出来
                title = list('td a ').text()
                url = list('td a ').attr('href')
                date = list('td:last-child').text()
                # 如需写入word，请将getmysql方法换成get_word即可
                getmysql(title,url,date)


        else:
            title = list('td a ').text()
            url = list('td a ').attr('href')
            date = list('td:last-child').text()
            getmysql(title, url, date)

# 上面完成单页信息的采集，现在进行前五页的信息采集。
def get_page():
    for i in range(1,6):

        print('开始抓取第%s页'%i)
        # 由于第一页不用点击操作我们从第二页开始进行点击操作
        if i > 1:
            brower.find_element_by_link_text(str(i)).click()
            # 这块停顿两秒，让页面内容顺利加载出来
            time.sleep(2)
        html = brower.page_source
        get_onepage(html)

        print('抓取第%s页完毕'%i)

    brower.close()

# 该函数将获取到的内容写入到word文件中

def get_word(title,url,date):

    document.add_paragraph(title)
    document.add_paragraph('网址：' + url)
    document.add_paragraph(date + '\n')

# 此函数将爬到的数据最近写到word中
def execute():
    # 给文档添加标题
    header = '{}{}招标项目清单'.format(shuru,shuru2)
    document.add_heading(header,level=0)
    # 运行爬虫程序
    get_page()
    # 将爬到的数据保存
    document.save('{}{}招标清单.docx'.format(shuru,shuru2))

# 该函数完成将数据写入mysql的操作
def getmysql(title,url,date):
    sql = "insert into shuju(title,url,riqi) values('%s','%s','%s')" %(title, url, date)
    # 执行sql语句
    cursor.execute(sql)


# 执行数据库写入操作
def main():
    # 执行get_page函数，将所有的数据写入到数据库
    print('开始执行爬虫')
    get_page()
    print('爬虫执行完毕，并关掉数据库')
    # 提交数据，关闭数据库
    conn.commit()
    conn.close()


main()

